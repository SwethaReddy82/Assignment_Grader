"""
create_sample_data.py
---------------------
Generates sample assignment instructions, rubric, and three student
submissions (DOCX format) so you can test the grading app immediately
without needing real student files.

Run:  python create_sample_data.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT = Path("sample_data")
SUBMISSIONS = OUTPUT / "student_submissions"
OUTPUT.mkdir(exist_ok=True)
SUBMISSIONS.mkdir(exist_ok=True)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


# ── Assignment Instructions ───────────────────────────────────────────────────

def create_instructions():
    doc = Document()
    add_heading(doc, "AI for Analytics — Assignment 1", level=1)
    add_heading(doc, "Exploratory Data Analysis with Python", level=2)

    add_para(doc, "Due Date: Week 4, Sunday 11:59 PM")
    add_para(doc, "Total Marks: 100 points")
    doc.add_paragraph()

    add_heading(doc, "Objective", level=3)
    add_para(doc,
        "In this assignment, you will perform a comprehensive Exploratory Data Analysis (EDA) "
        "on a real-world dataset of your choice. The goal is to demonstrate your ability to "
        "clean data, derive insights using statistical techniques, and communicate findings "
        "through visualisations and written analysis."
    )

    add_heading(doc, "Requirements", level=3)
    items = [
        "1. Dataset Selection: Choose a dataset from Kaggle, UCI ML Repository, or any public source. "
           "Justify your choice (2–3 sentences).",
        "2. Data Cleaning: Handle missing values, outliers, and data type issues. Document every step.",
        "3. Descriptive Statistics: Compute and interpret mean, median, standard deviation, skewness, "
           "and kurtosis for at least 5 numerical features.",
        "4. Visualisations: Create a minimum of 5 distinct visualisation types "
           "(histogram, box plot, scatter plot, heatmap, bar chart) using matplotlib or seaborn.",
        "5. Insights & Conclusions: Summarise 3–5 key findings supported by your analysis.",
        "6. Code Quality: Well-commented Python code (Jupyter Notebook or .py script). "
           "Use meaningful variable names and modular functions.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Submission Format", level=3)
    add_para(doc,
        "Submit a single PDF report (exported from Jupyter Notebook or Word) that includes "
        "your code snippets, outputs, visualisations, and written analysis. Late submissions "
        "will incur a 10% penalty per day."
    )

    path = OUTPUT / "assignment_instructions.docx"
    doc.save(str(path))
    print(f"  ✓  Created {path}")


# ── Grading Rubric ────────────────────────────────────────────────────────────

def create_rubric():
    doc = Document()
    add_heading(doc, "Grading Rubric — Assignment 1", level=1)
    add_para(doc, "Total: 100 Points")
    doc.add_paragraph()

    criteria = [
        ("Dataset Selection & Justification", 10,
         "Full marks: clear dataset description with strong justification and relevance to analytics. "
         "Partial: dataset described but justification weak. Zero: no justification provided."),
        ("Data Cleaning & Preprocessing", 20,
         "Full marks: all missing values, outliers, and type issues handled with documented rationale. "
         "Partial: some issues addressed but not all. Zero: no cleaning performed."),
        ("Descriptive Statistics", 20,
         "Full marks: all 5+ features analysed with correct stats and clear interpretation. "
         "Partial: fewer than 5 features or statistics computed but not interpreted. Zero: missing."),
        ("Visualisations", 25,
         "Full marks: 5+ distinct chart types, labelled, readable, and each accompanied by an insight. "
         "Partial: fewer than 5 types or charts present but poorly explained. Zero: no visuals."),
        ("Insights & Conclusions", 15,
         "Full marks: 3–5 specific, data-supported insights presented clearly. "
         "Partial: insights vague or fewer than 3. Zero: no conclusions."),
        ("Code Quality & Documentation", 10,
         "Full marks: code is modular, well-commented, and follows PEP 8. "
         "Partial: some comments but inconsistent. Zero: no documentation."),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Criterion"
    hdr[1].text = "Max Points"
    hdr[2].text = "Grading Notes"
    for name, pts, notes in criteria:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(pts)
        row[2].text = notes

    path = OUTPUT / "grading_rubric.docx"
    doc.save(str(path))
    print(f"  ✓  Created {path}")


# ── Student Submissions ───────────────────────────────────────────────────────

def create_student(filename, name, quality):
    """quality: 'excellent', 'average', 'poor'"""
    doc = Document()
    add_heading(doc, f"Assignment 1 — {name}", level=1)
    add_para(doc, f"Student: {name}  |  Course: AI for Analytics  |  Date: 2026-04-01")
    doc.add_paragraph()

    if quality == "excellent":
        add_heading(doc, "1. Dataset Selection", level=2)
        add_para(doc,
            "I selected the 'World Happiness Report 2023' dataset from Kaggle (156 countries, 9 features). "
            "This dataset is ideal for EDA because it combines economic, social, and health indicators, "
            "making it rich for multi-dimensional analysis relevant to AI for Analytics."
        )
        add_heading(doc, "2. Data Cleaning", level=2)
        add_para(doc,
            "The dataset had 3 missing values in 'Generosity' column — imputed using the column median. "
            "Two outliers in 'GDP per capita' were detected using IQR method (values > Q3 + 1.5*IQR) "
            "and retained since they represent real extreme countries. All columns were confirmed to "
            "have correct dtypes. Created a cleaned copy to preserve the raw data."
        )
        doc.add_paragraph("Code snippet:", style="Intense Quote")
        doc.add_paragraph(
            "df['Generosity'].fillna(df['Generosity'].median(), inplace=True)\n"
            "# IQR outlier detection\n"
            "Q1 = df['GDP'].quantile(0.25); Q3 = df['GDP'].quantile(0.75)\n"
            "IQR = Q3 - Q1"
        )
        add_heading(doc, "3. Descriptive Statistics", level=2)
        add_para(doc,
            "Computed statistics for 6 features: Happiness Score (mean=5.52, std=1.10, skew=-0.01), "
            "GDP per capita (mean=9.42, std=1.15, skew=-0.54), Social Support (mean=0.81, skew=-0.95), "
            "Healthy Life Expectancy (mean=64.0, std=6.8, kurtosis=0.63), Freedom (mean=0.73, skew=-1.0), "
            "Corruption Perception (mean=0.74, skew=2.1 — strongly right-skewed indicating most "
            "countries perceive low corruption). The skew in corruption suggests potential log-transform."
        )
        add_heading(doc, "4. Visualisations", level=2)
        add_para(doc,
            "Created 6 visualisations: (1) Histogram of Happiness Scores — near-normal distribution. "
            "(2) Box plot by region showing Europe > SE Asia > Africa. (3) Scatter plot: GDP vs "
            "Happiness (r=0.78, strong positive). (4) Correlation heatmap revealing GDP, social "
            "support, and life expectancy cluster together (r>0.7). (5) Bar chart: Top 10 happiest "
            "countries — all Nordic. (6) Pair plot showing non-linear relationship between freedom "
            "and happiness above threshold=0.6."
        )
        add_heading(doc, "5. Insights & Conclusions", level=2)
        add_para(doc,
            "Key findings: (1) GDP per capita is the strongest predictor of happiness (r=0.78). "
            "(2) Nordic countries dominate top rankings, suggesting cultural factors beyond economics. "
            "(3) Corruption perception has very low direct correlation with happiness (r=0.12) but "
            "clusters regionally. (4) Freedom shows a threshold effect — happiness plateaus for "
            "freedom scores below 0.5. (5) Life expectancy and GDP are co-linear (r=0.86), "
            "indicating potential multicollinearity in future modelling."
        )
        add_heading(doc, "6. Code Quality", level=2)
        add_para(doc,
            "All code written as modular functions (load_data(), clean_data(), compute_stats(), "
            "plot_visualisations()). PEP 8 compliant, with docstrings and inline comments on "
            "non-obvious operations. Used requirements.txt to pin library versions."
        )

    elif quality == "average":
        add_heading(doc, "Dataset", level=2)
        add_para(doc, "I used the Titanic dataset from Kaggle. It has passenger information.")
        add_heading(doc, "Data Cleaning", level=2)
        add_para(doc,
            "I dropped rows with missing values using df.dropna(). There were some null values "
            "in the Age column so those rows were removed."
        )
        add_heading(doc, "Statistics", level=2)
        add_para(doc,
            "Mean age: 29.7. Mean fare: 32.2. Survival rate: 38%. "
            "Standard deviation of age is 14.5."
        )
        add_heading(doc, "Visualisations", level=2)
        add_para(doc,
            "I made a histogram of ages, a bar chart of survival by gender, "
            "and a scatter plot of fare vs age. Also a heatmap of correlations. "
            "The charts show that women survived more than men."
        )
        add_heading(doc, "Conclusions", level=2)
        add_para(doc,
            "Women had a higher survival rate. Higher fares were linked to better survival. "
            "Age didn't seem to matter much."
        )

    else:  # poor
        add_heading(doc, "EDA Report", level=2)
        add_para(doc,
            "I downloaded a dataset about movies. It had columns like title, rating, and year."
        )
        add_heading(doc, "Analysis", level=2)
        add_para(doc,
            "The average rating was about 6.5. I made a bar chart but it was hard to read. "
            "I also looked at the data and it seemed normal."
        )
        add_heading(doc, "Conclusion", level=2)
        add_para(doc, "Movies have different ratings. Some are better than others.")

    path = SUBMISSIONS / filename
    doc.save(str(path))
    print(f"  ✓  Created {path}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("\nCreating sample data for AI Grading Assistant demo...\n")

    try:
        from docx import Document
    except ImportError:
        print("❌  python-docx not installed. Run: pip install python-docx")
        raise SystemExit(1)

    create_instructions()
    create_rubric()
    create_student("alice_chen_assignment1.docx",    "Alice Chen",    "excellent")
    create_student("bob_martinez_assignment1.docx",  "Bob Martinez",  "average")
    create_student("carol_liu_assignment1.docx",     "Carol Liu",     "poor")

    print(f"""
Done! Sample files created in ./{OUTPUT}/

To run the grader:
    python main.py \\
        --instructions sample_data/assignment_instructions.docx \\
        --rubric       sample_data/grading_rubric.docx \\
        --submissions  sample_data/student_submissions/ \\
        --output       sample_data/grading_output/ \\
        --assignment   "Assignment 1: EDA with Python"
""")
